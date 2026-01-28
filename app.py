import streamlit as st
import requests
import json
import base64
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import io
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from datetime import datetime
import time
from pdf2image import convert_from_bytes

# ==========================================
# 1. إعدادات الصفحة والتصميم
# ==========================================
st.set_page_config(page_title="Medical Notes Converter", page_icon="🩺", layout="centered")

# CSS بسيط لتحسين الشكل
st.markdown("""
    <style>
    .main { direction: rtl; }
    .stButton>button { width: 100%; border-radius: 10px; }
    h1 { color: #0e76a8; text-align: center; }
    </style>
    """, unsafe_allow_html=True)

st.title("🩺 المساعد الطبي الذكي")
st.markdown("---")

# ==========================================
# 2. الدوال الأساسية (Core Functions)
# ==========================================

# --- دالة تجهيز ملف الوورد (التنسيق الطبي) ---
def create_medical_doc():
    doc = Document()
    
    # ضبط الخط الأساسي: Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    # إجبار الخط للعربي والإنجليزي
    rPr = style.element.get_or_add_rPr()
    rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')

    # ضبط العناوين (Heading 1)
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_font.color.rgb = None # لون أسود

    # إضافة إطار للصفحة (Box Border)
    sections = doc.sections
    for section in sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12') # 1.5 pt
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
            
    return doc

# --- دالة الاتصال بـ Gemini (مباشرة وبسيطة) ---
def ask_gemini(api_key, image_bytes, mime_type="image/jpeg"):
    # نستخدم الموديل القياسي الثابت 1.5 Flash
    # ده الموديل اللي عليه عرض 1500 طلب مجاني يومياً
    MODEL_NAME = "gemini-1.5-flash"
    
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL_NAME}:generateContent?key={api_key}"
    
    try:
        # تحويل الصورة لكود Base64
        b64_image = base64.b64encode(image_bytes).decode('utf-8')
    except:
        return None, "فشل في قراءة ملف الصورة."

    headers = {'Content-Type': 'application/json'}
    
    # الأمر الموجه للذكاء الاصطناعي
    prompt = """
    You are a professional Medical Scribe.
    1. Transcribe the text from this medical image exactly.
    2. Format HEADINGS by starting the line with # (e.g., # Diagnosis).
    3. Keep body text as normal paragraphs.
    4. Do not use Markdown bold (**) or italics (*).
    """
    
    payload = {
        "contents": [{
            "parts": [
                {"text": prompt},
                {"inline_data": {"mime_type": mime_type, "data": b64_image}}
            ]
        }],
        "safetySettings": [
             {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
             {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
             {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
             {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
        ]
    }
    
    # المحاولة مرة واحدة (سنعتمد على التأخير الزمني لمنع الخطأ أصلاً)
    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        
        if response.status_code == 200:
            return response.json()['candidates'][0]['content']['parts'][0]['text'], None
        elif response.status_code == 429:
            return None, "خطأ ضغط (429): تجاوزنا الحد المسموح في الدقيقة."
        elif response.status_code == 404:
            return None, f"خطأ موديل (404): الموديل {MODEL_NAME} غير متاح لهذا المفتاح."
        else:
            return None, f"خطأ غير معروف ({response.status_code})"
            
    except Exception as e:
        return None, f"خطأ اتصال: {str(e)}"

# --- دالة الفيدباك (Google Sheets) ---
def save_feedback(text):
    if "gcp_service_account" in st.secrets:
        try:
            scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
            creds_dict = dict(st.secrets["gcp_service_account"])
            creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
            client = gspread.authorize(creds)
            sheet = client.open("Medical_App_Feedback").sheet1 
            sheet.append_row([datetime.now().strftime("%Y-%m-%d %H:%M:%S"), text])
            return True
        except: return False
    return False

# ==========================================
# 3. واجهة التطبيق (Streamlit App)
# ==========================================

# التحقق من المفتاح
if "GEMINI_API_KEY" not in st.secrets:
    st.error("⚠️ ملف الأسرار (Secrets) غير موجود أو لا يحتوي على GEMINI_API_KEY.")
    st.stop()

api_key = st.secrets["GEMINI_API_KEY"]

# مدخلات المستخدم
col1, col2 = st.columns([2, 1])
with col1:
    doc_title = st.text_input("اسم ملف الوورد:", value="Medical Notes")
with col2:
    st.write("")
    st.write("")
    hide_names = st.checkbox("إخفاء أسماء الصور؟", value=False)

uploaded_files = st.file_uploader("ارفع الصور أو ملف PDF", type=["pdf", "jpg", "jpeg", "png"], accept_multiple_files=True)

# زر البدء
if uploaded_files and st.button("🚀 ابدأ التحويل الآن"):
    
    # تجهيز ملف الوورد
    doc = create_medical_doc()
    # إضافة العنوان الرئيسي
    title_para = doc.add_paragraph(doc_title, style='Title')
    title_para.alignment = 1 # Center
    
    # متغيرات التتبع
    progress_bar = st.progress(0)
    status_text = st.empty()
    files_processed = 0
    
    # --- تحضير القائمة الكاملة للصور (فك الـ PDF لو موجود) ---
    all_images_to_process = [] # قائمة فيها (الصورة، الاسم)
    
    with st.spinner("جاري قراءة الملفات..."):
        for file in uploaded_files:
            if file.type == "application/pdf":
                try:
                    # تحويل PDF لصور
                    pdf_pages = convert_from_bytes(file.read())
                    for i, page in enumerate(pdf_pages):
                        # تحويل لـ Bytes
                        img_byte_arr = io.BytesIO()
                        page.save(img_byte_arr, format='JPEG')
                        all_images_to_process.append({
                            "bytes": img_byte_arr.getvalue(),
                            "name": f"{file.name} (صفحة {i+1})",
                            "type": "image/jpeg"
                        })
                except Exception as e:
                    st.error(f"مشكلة في ملف PDF: {file.name}")
            else:
                # صورة عادية
                all_images_to_process.append({
                    "bytes": file.getvalue(),
                    "name": file.name,
                    "type": file.type
                })

    total_count = len(all_images_to_process)
    
    # --- بداية المعالجة الفعلية ---
    for i, item in enumerate(all_images_to_process):
        current_step = i + 1
        status_text.write(f"⏳ جاري معالجة {current_step}/{total_count}: **{item['name']}**...")
        
        # 1. إرسال لـ Gemini
        text, error = ask_gemini(api_key, item['bytes'], item['type'])
        
        if error:
            st.error(f"خطأ في {item['name']}: {error}")
            doc.add_paragraph(f"[فشل استخراج النص من: {item['name']} - السبب: {error}]")
        else:
            # 2. الكتابة في الوورد
            if not hide_names:
                doc.add_heading(item['name'], level=1)
            
            # تنسيق النص (تحويل # لعناوين)
            for line in text.split('\n'):
                line = line.strip()
                if not line: continue
                if line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), level=1)
                else:
                    doc.add_paragraph(line)
            
            doc.add_page_break()
        
        # تحديث الشريط
        progress_bar.progress(current_step / total_count)
        
        # 3. استراحة إجبارية (4 ثواني) لتجنب Error 429
        # لا تستنا في آخر صورة
        if current_step < total_count:
            time.sleep(4) 

    status_text.success("✅ تم الانتهاء بنجاح!")
    
    # زر التحميل
    bio = io.BytesIO()
    doc.save(bio)
    st.download_button(
        label="📥 تحميل الملف (Word)",
        data=bio.getvalue(),
        file_name=f"{doc_title}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary"
    )

# الفيدباك
st.markdown("---")
with st.expander("💬 رأيك يهمنا"):
    with st.form("fb_form"):
        txt = st.text_area("اكتب ملاحظاتك هنا:")
        if st.form_submit_button("إرسال"):
            if save_feedback(txt): st.success("وصلنا، شكراً ليك!")
            else: st.error("خطأ في الإرسال")
